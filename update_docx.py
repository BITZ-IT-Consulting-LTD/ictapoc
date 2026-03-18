from docx import Document

doc_path = "mdas/docs_final/priority_mdas/Prioritisation Framework for Government Digitisation  Automation - v03 - Approved - Priority MDAs.docx"
doc = Document(doc_path)
table = doc.tables[0]

# 2. Agriculture change to Livestock
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if "Agriculture" in run.text:
                    run.text = run.text.replace("Agriculture", "Livestock")

# Check if Science, Research & Innovation exists
science_exists = any("Science" in row.cells[1].text for row in table.rows[1:])

new_mdas = [
    "TVET",
    "Internal Security and National Administration",
    "Cabinet Office",
    "State Dept for Co-operatives",
    "Office of the Attorney General — Marriages",
    "State Dept for Parliamentary Affairs"
]

if not science_exists:
    new_mdas.append("State Dept for Science Research and Innovation")

# We want the ID to follow the previous number.
# The last row id is in table.rows[-1].cells[0].text
try:
    current_num = int(table.rows[-1].cells[0].text.strip()) + 1
except ValueError:
    current_num = len(table.rows)

for mda in new_mdas:
    row = table.add_row()
    row.cells[0].text = str(current_num)
    row.cells[1].text = mda
    # Fill remaining columns with some placeholder or empty string to match structure
    for i in range(2, 9):
        row.cells[i].text = "-"
    current_num += 1

doc.save(doc_path)
print("Document updated successfully.")
